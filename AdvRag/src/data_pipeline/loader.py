"""
data_pipeline/loader.py
-----------------------
Pure document loading layer — reads files and returns raw LlamaIndex Document
objects with layout-aware metadata.

Responsibility boundary (important)
-------------------------------------
This file ONLY reads and enriches metadata. It does NOT split into chunks.
Chunking, LLM metadata extraction, and de-duplication are all handled by
transformation.py (IngestionPipeline). Keeping these separated means:
  • loader.py can be tested without a live LLM.
  • transformation.py can receive Documents from any source (loader, LlamaParse,
    helpers.py, a REST API) without caring how they were loaded.

Supported file types
---------------------
.pdf   → PyMuPDFReader  (fast, layout-aware, no cloud API needed)
.docx  → DocxReader
.md    → MarkdownReader
.txt   → TextReader

For complex PDFs / DOCX with tables and multi-column layouts, use
utils/helpers.py::aparse_with_llamaparse() instead.
"""

import re
import uuid
import asyncio
from pathlib import Path
from typing import List
from docx import Document as DocxDocument

from llama_index.core import Document as LlamaDocument
from llama_index.readers.file import (
    PyMuPDFReader,
    DocxReader,
    MarkdownReader,
    FlatReader,
)

# Single, correct import for the installed package
from config import settings  # noqa: F401 (imported for consistency; not used directly here)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(
    r"^(#{1,6}\s+.+|[A-Z][A-Z0-9 \-:]{3,}(?:\n[=\-]{3,})?)",
    re.MULTILINE,
)


# ---------------------------------------------------------------------------
# Header extraction
# ---------------------------------------------------------------------------

def _extract_first_header(text: str) -> str:
    """
    Extracts the first Markdown or ALL-CAPS heading found in a text block.
    Returns 'Document Start' if none is found.
    """
    match = _HEADING_RE.search(text)
    if match:
        return match.group(0).lstrip("#").strip().splitlines()[0]
    return "Document Start"


# ---------------------------------------------------------------------------
# Reader factory
# ---------------------------------------------------------------------------

def _get_reader(suffix: str):
    """Maps file extensions to LlamaIndex reader instances."""
    return {
        ".pdf":  PyMuPDFReader(),
        ".docx": DocxReader(),
        ".md":   MarkdownReader(),
        ".txt":  FlatReader(),
    }.get(suffix)



# In loader.py, replace DocxReader with this custom function:
from docx import Document as DocxDocument

def load_docx_with_tables(file_path: Path) -> List[LlamaDocument]:
    doc = DocxDocument(file_path)
    docs = []
    
    # Extract paragraphs as one document
    para_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if para_text:
        docs.append(LlamaDocument(text=para_text, metadata={"content_type": "text"}))
    
    # Each table becomes its own Document — never split
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        table_text = "\n".join(rows)
        if table_text.strip():
            docs.append(LlamaDocument(
                text=table_text,
                metadata={"content_type": "table", "table_index": i}
            ))
    
    return docs
# ---------------------------------------------------------------------------
# Core loader  —  returns raw Documents, NOT split nodes
# ---------------------------------------------------------------------------

async def aload_documents_from_directory(
    directory: str | Path,
) -> List[LlamaDocument]:
    """
    Recursively loads all supported files from `directory` and returns a flat
    list of LlamaIndex Document objects with enriched metadata.

    Does NOT split documents into chunks — pass the output to
    transformation.arun_pipeline() for chunking + metadata extraction.

    Metadata injected on every Document
    ------------------------------------
    file_name       : basename of the source file
    file_path       : absolute path string
    file_type       : extension without the leading dot  ("pdf", "md", …)
    section_header  : first heading found in the document / page text
    doc_id          : stable UUID tied to the file stem (for de-duplication)

    Args:
        directory: Root directory to scan (scanned recursively).

    Returns:
        List of Document objects ready for the IngestionPipeline.
    """
    directory = Path(directory)
    all_docs: List[LlamaDocument] = []

    for file_path in sorted(directory.rglob("*")):
        suffix = file_path.suffix.lower()
        reader = _get_reader(suffix)
        if reader is None:
            continue

        print(f"[Loader] Reading: {file_path.name}")

        try:
            # Readers are synchronous — offload to thread pool.
            loop = asyncio.get_running_loop()
            raw_docs: List[LlamaDocument] = await loop.run_in_executor(
                None, reader.load_data, file_path
            )

            current_header = "Document Start"
            for doc in raw_docs:
                found = _extract_first_header(doc.get_content())
                if found != "Document Start":
                    current_header = found

                doc.metadata.update(
                    {
                        "file_name":      file_path.name,
                        "file_path":      str(file_path),
                        "file_type":      suffix.lstrip("."),
                        "section_header": current_header,
                        "doc_id":         (
                            f"{file_path.stem}-"
                            f"{uuid.uuid5(uuid.NAMESPACE_URL, str(file_path)).hex[:12]}"
                        ),
                    }
                )

            all_docs.extend(raw_docs)
            print(f"[Loader] → {len(raw_docs)} document(s) from '{file_path.name}'")

        except Exception as exc:
            print(f"[Loader] ERROR reading '{file_path.name}': {exc}")

    print(f"[Loader] Total documents loaded: {len(all_docs)}")
    return all_docs


async def aload_single_file(file_path: str | Path) -> List[LlamaDocument]:
    """
    Loads a single file and returns its Document(s).
    Convenience wrapper around aload_documents_from_directory for upload
    endpoints that receive one file at a time.

    Args:
        file_path: Absolute or relative path to a supported file.

    Returns:
        List of Document objects (multiple for multi-page PDFs).

    Raises:
        ValueError: If the file type is not supported.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    reader = _get_reader(suffix)

    if reader is None:
        raise ValueError(
            f"Unsupported file type: '{suffix}'. "
            f"Supported: .pdf, .docx, .md, .txt"
        )

    print(f"[Loader] Reading single file: {path.name}")
    loop = asyncio.get_running_loop()
    raw_docs: List[LlamaDocument] = await loop.run_in_executor(
        None, reader.load_data, path
    )

    current_header = "Document Start"
    for doc in raw_docs:
        found = _extract_first_header(doc.get_content())
        if found != "Document Start":
            current_header = found
        doc.metadata.update(
            {
                "file_name":      path.name,
                "file_path":      str(path),
                "file_type":      suffix.lstrip("."),
                "section_header": current_header,
                "doc_id":         (
                    f"{path.stem}-"
                    f"{uuid.uuid5(uuid.NAMESPACE_URL, str(path)).hex[:12]}"
                ),
            }
        )

    return raw_docs


# ---------------------------------------------------------------------------
# Utility: display
# ---------------------------------------------------------------------------

def summarise_documents(docs: List[LlamaDocument], n: int = 3) -> None:
    """Pretty-prints the first `n` Documents for debugging."""
    for i, doc in enumerate(docs[:n]):
        meta = doc.metadata
        print(
            f"\n--- Document {i} ---\n"
            f"  file_name:      {meta.get('file_name')}\n"
            f"  file_type:      {meta.get('file_type')}\n"
            f"  section_header: {meta.get('section_header')}\n"
            f"  doc_id:         {meta.get('doc_id')}\n"
            f"  content[:120]:  {doc.get_content()[:120]!r}"
        )


# ---------------------------------------------------------------------------
# Smoke test
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import tempfile

    async def _test():
        with tempfile.TemporaryDirectory() as tmpdir:
            test_file = Path(tmpdir) / "sample.md"
            test_file.write_text(
                "# Experience\n\nBuilt RAG systems with Groq + Mixedbread.\n\n"
                "# Education\n\nB.Sc. Computer Science.",
                encoding="utf-8",
            )

            docs = await aload_documents_from_directory(tmpdir)
            summarise_documents(docs)
            print(
                f"\n[Smoke] Loaded {len(docs)} document(s). "
                f"Pass to transformation.arun_pipeline() for chunking."
            )

    asyncio.run(_test())